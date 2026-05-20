"""问答型知识库切块器。

这个文件把 Excel、CSV/TXT、JSON、Markdown、Docx、PDF 中的问答对
转换成统一 chunk 结构，是医疗 QA 知识库建库时的重要入口。

支持的格式：
- Excel（.xlsx/.xls）：默认前两列分别是问题和答案
- CSV/TXT：UTF-8，TAB 或逗号分隔，两列问答
- JSON/JSONL：支持 {"question": ..., "answer": ...} 或 {"q": ..., "a": ...}
- PDF/Markdown/Docx：基于标题、列表或问答结构提取

每一组问答对都会被视为一个原子 chunk。
"""

#

import json
import logging
import re
import csv
from copy import deepcopy
from io import BytesIO
from timeit import default_timer as timer
from openpyxl import load_workbook

from parser import get_text
from rag.nlp import is_english, random_choices, qbullets_category, add_positions, has_qbullet, docx_question_level
from rag.nlp import rag_tokenizer, tokenize_table, concat_img
from markdown import markdown

from common.float_utils import get_float


class Excel:
    """Excel 文件问答对提取器。"""
    def __call__(self, fnm, binary=None, callback=None):
        """从 Excel 中提取问答对，默认按前两列读取为问和答。
        
        Args:
            fnm: 文件名
            binary: 文件二进制内容（可选）
            callback: 进度回调函数
            
        Returns:
            list: 问答对列表 [(q1, a1), (q2, a2), ...]
        """
        if not binary:
            wb = load_workbook(fnm)
        else:
            wb = load_workbook(BytesIO(binary))
        
        # 计算总行数用于进度计算
        total = 0
        for sheetname in wb.sheetnames:
            total += len(list(wb[sheetname].rows))

        res, fails = [], []
        for sheetname in wb.sheetnames:
            ws = wb[sheetname]
            rows = list(ws.rows)
            for i, r in enumerate(rows):
                q, a = "", ""
                # 前两列分别作为问题和答案
                for cell in r:
                    if not cell.value:
                        continue
                    if not q:
                        q = str(cell.value)
                    elif not a:
                        a = str(cell.value)
                    else:
                        break
                if q and a:
                    res.append((q, a))
                else:
                    fails.append(str(i + 1))
                # 每999条更新一次进度
                if len(res) % 999 == 0:
                    callback(len(res) * 0.6 / total, 
                             f"Extract pairs: {len(res)}" + 
                             (f" {len(fails)} failure, line: {",".join(fails[:3])}..." if fails else ""))

        callback(0.6, f"Extract pairs: {len(res)}. " + 
                 (f"{len(fails)} failure, line: {",".join(fails[:3])}..." if fails else ""))
        
        # 检测语言
        self.is_english = is_english(
            [rmPrefix(q) for q, _ in random_choices(res, k=30) if len(q) > 1])
        return res



class Docx:
    """Docx 文件问答对提取器。
    
    基于标题层级和正文内容抽取问答对，支持多级标题嵌套。
    """
    def __init__(self):
        pass

    def __call__(self, filename, binary=None, from_page=0, to_page=100000, callback=None):
        """从 Docx 的标题层级和正文内容中抽取问答对。
        
        Args:
            filename: 文件名
            binary: 文件二进制内容（可选）
            from_page: 起始页码
            to_page: 结束页码
            callback: 进度回调函数
            
        Returns:
            tuple: (问答对列表, 表格列表)
        """
        from docx import Document
        self.doc = Document(
            filename) if not binary else Document(BytesIO(binary))
        
        pn = 0  # 当前页码
        last_answer, last_image = "", None
        question_stack, level_stack = [], []  # 问题栈和层级栈
        qai_list = []  # 问答对列表
        
        for p in self.doc.paragraphs:
            if pn > to_page:
                break
            question_level, p_text = 0, ''
            if from_page <= pn < to_page and p.text.strip():
                question_level, p_text = docx_question_level(p)
            
            # 不是问题标题，视为答案正文
            if not question_level or question_level > 6:
                last_answer = f'{last_answer}\n{p_text}'
                current_image = self.get_picture(self.doc, p)
                last_image = concat_img(last_image, current_image)
            else:
                # 命中了问题标题层级，保存上一个问答对
                if last_answer or last_image:
                    sum_question = '\n'.join(question_stack)
                    if sum_question:
                        qai_list.append((sum_question, last_answer, last_image))
                    last_answer, last_image = '', None

                # 更新问题栈（处理嵌套标题）
                i = question_level
                while question_stack and i <= level_stack[-1]:
                    question_stack.pop()
                    level_stack.pop()
                question_stack.append(p_text)
                level_stack.append(question_level)
            
            # 检测分页
            for run in p.runs:
                if 'lastRenderedPageBreak' in run._element.xml:
                    pn += 1
                    continue
                if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                    pn += 1
        
        # 处理最后一个问答对
        if last_answer:
            sum_question = '\n'.join(question_stack)
            if sum_question:
                qai_list.append((sum_question, last_answer, last_image))

        # 提取表格
        tbls = []
        for tb in self.doc.tables:
            html = "<table>"
            for r in tb.rows:
                html += "<tr>"
                i = 0
                while i < len(r.cells):
                    span = 1
                    c = r.cells[i]
                    for j in range(i + 1, len(r.cells)):
                        if c.text == r.cells[j].text:
                            span += 1
                            i = j
                    i += 1
                    html += f"<td>{c.text}</td>" if span == 1 else f"<td colspan='{span}'>{c.text}</td>"
                html += "</tr>"
            html += "</table>"
            tbls.append(((None, html), ""))
        
        return qai_list, tbls


def rmPrefix(txt):
    """移除文本开头的问答前缀（如"问题："、"Q:"等）。
    
    Args:
        txt: 原始文本
        
    Returns:
        str: 移除前缀后的文本
    """
    return re.sub(
        r"^(问题|答案|回答|user|assistant|Q|A|Question|Answer|问|答)[\t:： ]+", "", txt.strip(), flags=re.IGNORECASE)



def beAdocDocx(d, q, a, eng, image, row_num=-1):
    """将问答对转换为带图片的chunk文档结构。
    
    Args:
        d: 基础文档字典
        q: 问题
        a: 答案
        eng: 是否为英文
        image: 图片内容
        row_num: 行号（用于排序）
        
    Returns:
        dict: 完整的chunk文档
    """
    qprefix = "Question: " if eng else "问题："
    aprefix = "Answer: " if eng else "回答："
    d["content_with_weight"] = "\t".join(
        [qprefix + rmPrefix(q), aprefix + rmPrefix(a)])
    d["content_ltks"] = rag_tokenizer.tokenize(q)
    d["content_sm_ltks"] = rag_tokenizer.fine_grained_tokenize(d["content_ltks"])
    if image:
        d["image"] = image
        d["doc_type_kwd"] = "image"
    if row_num >= 0:
        d["top_int"] = [row_num]
    return d


def beAdoc(d, q, a, eng, row_num=-1):
    """将问答对转换为chunk文档结构。
    
    Args:
        d: 基础文档字典
        q: 问题
        a: 答案
        eng: 是否为英文
        row_num: 行号（用于排序）
        
    Returns:
        dict: 完整的chunk文档
    """
    qprefix = "Question: " if eng else "问题："
    aprefix = "Answer: " if eng else "回答："
    d["content_with_weight"] = "\t".join(
        [qprefix + rmPrefix(q), aprefix + rmPrefix(a)])
    d["content_ltks"] = rag_tokenizer.tokenize(q)
    d["content_sm_ltks"] = rag_tokenizer.fine_grained_tokenize(d["content_ltks"])
    if row_num >= 0:
        d["top_int"] = [row_num]
    return d


def mdQuestionLevel(s):
    """解析 Markdown 标题层级。
    
    Args:
        s: Markdown 文本行
        
    Returns:
        tuple: (标题级别, 标题文本)
    """
    match = re.match(r'#*', s)
    return (len(match.group(0)), s.lstrip('#').lstrip()) if match else (0, s)


def chunk(filename, binary=None, from_page=0, to_page=100000, lang="Chinese", callback=None, **kwargs):
    """将各种格式的文档转换为问答型chunk。
    
    支持格式：
    - Excel（.xlsx/.xls）：默认前两列分别是问题和答案，无表头要求
    - CSV/TXT：UTF-8，TAB 或逗号分隔，两列问答
    - JSON/JSONL：支持 {"question": ..., "answer": ...} 或 {"q": ..., "a": ...}
    - PDF/Markdown/Docx：基于标题、列表或问答结构提取

    每一组问答对都会被视为一个原子 chunk。
    无法解析的行/记录会被跳过，并通过回调上报。
    
    Args:
        filename: 文件名（用于判断格式）
        binary: 文件二进制内容
        from_page: 起始页码
        to_page: 结束页码
        lang: 语言（Chinese/English）
        callback: 进度回调函数
        
    Returns:
        list: chunk文档列表
    """
    eng = lang.lower() == "english"
    res = []
    doc = {
        "docnm_kwd": filename,
        "title_tks": rag_tokenizer.tokenize(re.sub(r"\.[a-zA-Z]+$", "", filename))
    }
    
    # Excel格式处理
    if re.search(r"\.xlsx?$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        excel_parser = Excel()
        for ii, (q, a) in enumerate(excel_parser(filename, binary, callback)):
            res.append(beAdoc(deepcopy(doc), q, a, eng, ii))
        return res

    elif re.search(r"\.(txt)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        txt = get_text(filename, binary)
        lines = txt.split("\n")
        comma, tab = 0, 0
        for line in lines:
            if len(line.split(",")) == 2:
                comma += 1
            if len(line.split("\t")) == 2:
                tab += 1
        delimiter = "\t" if tab >= comma else ","

        fails = []
        question, answer = "", ""
        i = 0
        while i < len(lines):
            arr = lines[i].split(delimiter)
            if len(arr) != 2:
                if question:
                    answer += "\n" + lines[i]
                else:
                    fails.append(str(i + 1))
            elif len(arr) == 2:
                if question and answer:
                    res.append(beAdoc(deepcopy(doc), question, answer, eng, i))
                question, answer = arr
            i += 1
            if len(res) % 999 == 0:
                callback(len(res) * 0.6 / len(lines), ("Extract Q&A: {}".format(len(res)) + (
                    f"{len(fails)} failure, line: %s..." % (",".join(fails[:3])) if fails else "")))

        if question:
            res.append(beAdoc(deepcopy(doc), question, answer, eng, len(lines)))

        callback(0.6, ("Extract Q&A: {}".format(len(res)) + (
            f"{len(fails)} failure, line: %s..." % (",".join(fails[:3])) if fails else "")))

        return res

    elif re.search(r"\.(csv)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        txt = get_text(filename, binary)
        lines = txt.split("\n")
        delimiter = "\t" if any("\t" in line for line in lines) else ","

        fails = []
        question, answer = "", ""
        res = []
        reader = csv.reader(lines, delimiter=delimiter)

        for i, row in enumerate(reader):
            if len(row) != 2:
                if question:
                    answer += "\n" + lines[i]
                else:
                    fails.append(str(i + 1))
            elif len(row) == 2:
                if question and answer:
                    res.append(beAdoc(deepcopy(doc), question, answer, eng, i))
                question, answer = row
            if len(res) % 999 == 0:
                callback(len(res) * 0.6 / len(lines), ("Extract Q&A: {}".format(len(res)) + (
                    f"{len(fails)} failure, line: %s..." % (",".join(fails[:3])) if fails else "")))

        if question:
            res.append(beAdoc(deepcopy(doc), question, answer, eng, len(list(reader))))

        callback(0.6, ("Extract Q&A: {}".format(len(res)) + (
            f"{len(fails)} failure, line: %s..." % (",".join(fails[:3])) if fails else "")))
        return res

    elif re.search(r"\.(md|markdown|mdx)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        txt = get_text(filename, binary)
        lines = txt.split("\n")
        _last_question, last_answer = "", ""
        question_stack, level_stack = [], []
        code_block = False
        for index, line in enumerate(lines):
            if line.strip().startswith('```'):
                code_block = not code_block
            question_level, question = 0, ''
            if not code_block:
                question_level, question = mdQuestionLevel(line)

            if not question_level or question_level > 6:  # not a question
                last_answer = f'{last_answer}\n{line}'
            else:  # is a question
                if last_answer.strip():
                    sum_question = '\n'.join(question_stack)
                    if sum_question:
                        res.append(beAdoc(deepcopy(doc), sum_question,
                                          markdown(last_answer, extensions=['markdown.extensions.tables']), eng, index))
                    last_answer = ''

                i = question_level
                while question_stack and i <= level_stack[-1]:
                    question_stack.pop()
                    level_stack.pop()
                question_stack.append(question)
                level_stack.append(question_level)
        if last_answer.strip():
            sum_question = '\n'.join(question_stack)
            if sum_question:
                res.append(beAdoc(deepcopy(doc), sum_question,
                                  markdown(last_answer, extensions=['markdown.extensions.tables']), eng, index))
        return res

    elif re.search(r"\.docx$", filename, re.IGNORECASE):
        docx_parser = Docx()
        qai_list, tbls = docx_parser(filename, binary,
                                     from_page=0, to_page=10000, callback=callback)
        res = tokenize_table(tbls, doc, eng)
        for i, (q, a, image) in enumerate(qai_list):
            res.append(beAdocDocx(deepcopy(doc), q, a, eng, image, i))
        return res

    elif re.search(r"\.(json|jsonl)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        txt = get_text(filename, binary)
        lines = txt.strip().splitlines() if re.search(r"\.jsonl$", filename, re.IGNORECASE) else None

        if lines is None:
        # JSON 数组格式。
            try:
                records = json.loads(txt)
                if isinstance(records, dict):
                    records = [records]
            except json.JSONDecodeError as ex:
                raise ValueError(f"Invalid JSON file: {ex}")
        else:
        # JSONL 格式：每行一条记录。
            records = []
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                try:
                    records.append(json.loads(line))
                except json.JSONDecodeError:
                    pass

        fails = []
        for i, record in enumerate(records):
            q = record.get("question") or record.get("q") or record.get("问题") or ""
            a = record.get("answer") or record.get("a") or record.get("回答") or record.get("答案") or ""
            q, a = str(q).strip(), str(a).strip()
            if q and a:
                res.append(beAdoc(deepcopy(doc), q, a, eng, i))
            else:
                fails.append(str(i + 1))
            if len(res) % 999 == 0:
                callback(len(res) * 0.6 / max(len(records), 1),
                         "Extract Q&A: {}".format(len(res)) +
                         (f", {len(fails)} failures" if fails else ""))

        callback(0.6, "Extract Q&A: {}".format(len(res)) +
                 (f", {len(fails)} failure(s) at record(s): {','.join(fails[:5])}" if fails else ""))
        return res

    raise NotImplementedError(
        "Supported formats: Excel, CSV/TXT, JSON, JSONL, PDF, Markdown, Docx.")


if __name__ == "__main__":
    import sys


    def dummy(prog=None, msg=""):
        pass


    chunk(sys.argv[1], from_page=0, to_page=10, callback=dummy)